#!/usr/bin/env python3
"""
WifeyMOOC JSON to Paper Exercise Printer
Converts exercise JSON files into printable HTML + DOCX format
"""

import json
import os
import sys
import base64
import random
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Tuple

# Try to import python-docx, but make it optional
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

class ExerciseToPaper:
    def __init__(self, json_file: str):
        self.json_file = json_file
        self.exercises = []
        self.base_dir = Path(json_file).parent
        self.full_page_images = {}  # Track images for full-page references
        self.load_json()
        
    def load_json(self):
        """Load and parse the JSON file"""
        try:
            with open(self.json_file, 'r', encoding='utf-8') as f:
                self.exercises = json.load(f)
            print(f"✓ Loaded {len(self.exercises)} exercises from {self.json_file}")
        except FileNotFoundError:
            print(f"✗ Error: File '{self.json_file}' not found")
            sys.exit(1)
        except json.JSONDecodeError as e:
            print(f"✗ Error: Invalid JSON - {e}")
            sys.exit(1)
    
    def _load_image_as_base64(self, image_path: str) -> str:
        """Load an image file and convert to base64 data URI"""
        try:
            # Try to find the image relative to the JSON file's directory
            full_path = self.base_dir / image_path
            if not full_path.exists():
                # Try relative to current directory
                full_path = Path(image_path)
            
            if not full_path.exists():
                print(f"⚠️  Warning: Image not found - {image_path}")
                return None
            
            # Determine MIME type
            suffix = full_path.suffix.lower()
            mime_types = {
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png',
                '.gif': 'image/gif',
                '.webp': 'image/webp',
                '.svg': 'image/svg+xml'
            }
            mime_type = mime_types.get(suffix, 'image/jpeg')
            
            # Read and encode image
            with open(full_path, 'rb') as f:
                image_data = base64.b64encode(f.read()).decode('utf-8')
            
            return f"data:{mime_type};base64,{image_data}"
        except Exception as e:
            print(f"⚠️  Warning: Could not load image {image_path} - {e}")
            return None
    
    def _get_image_path(self, image_path: str) -> Path:
        """Get the actual path to an image file"""
        full_path = self.base_dir / image_path
        if not full_path.exists():
            full_path = Path(image_path)
        return full_path if full_path.exists() else None
    
    def generate_html(self, output_file: str = None) -> str:
        """Generate HTML document for printing"""
        if output_file is None:
            output_file = Path(self.json_file).stem + "_paper.html"
        
        html_content = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>WifeyMOOC - Exercise Worksheet</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        html {
            font-size: 14px;
        }
        
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.5;
            color: #333;
            background: white;
            padding: 20px;
        }
        
        .container {
            max-width: 900px;
            margin: 0 auto;
            background: white;
        }
        
        .header {
            text-align: center;
            margin-bottom: 30px;
            border-bottom: 3px solid #2c3e50;
            padding-bottom: 15px;
            break-after: avoid;
        }
        
        .header h1 {
            color: #2c3e50;
            margin-bottom: 5px;
            font-size: 2em;
        }
        
        .header p {
            color: #7f8c8d;
            font-size: 0.9em;
        }
        
        .exercise {
            margin-bottom: 25px;
            border-left: 4px solid #3498db;
            padding-left: 15px;
            break-inside: avoid;
        }
        
        .exercise-number {
            display: inline-block;
            background: #3498db;
            color: white;
            padding: 4px 10px;
            border-radius: 15px;
            font-weight: bold;
            margin-bottom: 8px;
            font-size: 0.85em;
        }
        
        .exercise-type {
            display: inline-block;
            background: #ecf0f1;
            color: #2c3e50;
            padding: 3px 8px;
            border-radius: 4px;
            font-size: 0.75em;
            margin-left: 8px;
            font-weight: 600;
        }
        
        .question {
            font-size: 1.05em;
            margin: 12px 0;
            font-weight: 500;
            color: #2c3e50;
        }
        
        .media-note {
            background: #fff3cd;
            border-left: 3px solid #ffc107;
            padding: 8px 12px;
            margin: 10px 0;
            font-size: 0.85em;
            color: #856404;
        }
        
        .image-reference {
            background: #e8f4f8;
            border-left: 3px solid #3498db;
            padding: 8px 12px;
            margin: 10px 0;
            font-size: 0.85em;
            color: #2c3e50;
            font-style: italic;
        }
        
        .media-image {
            margin: 12px 0;
            text-align: center;
            border: 1px solid #dee2e6;
            padding: 10px;
            border-radius: 4px;
            background: #f8f9fa;
            break-inside: avoid;
        }
        
        .media-image img {
            max-width: 100%;
            height: auto;
            border-radius: 4px;
            display: inline-block;
            max-height: 250px;
        }
        
        .media-image-caption {
            font-size: 0.8em;
            color: #7f8c8d;
            margin-top: 6px;
            font-style: italic;
        }
        
        .image-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(130px, 1fr));
            gap: 12px;
            margin: 12px 0;
            padding: 12px;
            background: #f8f9fa;
            border-radius: 4px;
            break-inside: avoid;
        }
        
        .image-item {
            border: 1px solid #dee2e6;
            padding: 10px;
            border-radius: 4px;
            text-align: center;
            background: white;
            break-inside: avoid;
        }
        
        .image-item img {
            max-width: 100%;
            height: auto;
            max-height: 120px;
            margin-bottom: 8px;
        }
        
        .image-label {
            font-weight: bold;
            color: #2c3e50;
            font-size: 0.9em;
        }
        
        .option-list {
            margin: 10px 0;
            padding-left: 15px;
        }
        
        .option {
            margin: 8px 0;
            padding: 8px;
            background: #f8f9fa;
            border-radius: 3px;
            border-left: 2px solid #dee2e6;
            font-size: 0.95em;
        }
        
        .option input[type="checkbox"],
        .option input[type="radio"] {
            margin-right: 8px;
            cursor: pointer;
            accent-color: #3498db;
        }
        
        .option label {
            cursor: pointer;
            user-select: none;
        }
        
        .sentence-parts {
            background: #f8f9fa;
            padding: 12px;
            border-radius: 4px;
            margin: 10px 0;
            line-height: 1.8;
            font-size: 0.95em;
            break-inside: avoid;
        }
        
        .answer-space {
            border-bottom: 1.5px solid #2c3e50;
            display: inline-block;
            min-width: 120px;
            margin: 0 4px;
            height: 18px;
        }
        
        .pairs-list {
            margin: 10px 0;
            padding: 12px;
            background: #f8f9fa;
            border-radius: 4px;
            break-inside: avoid;
        }
        
        .pair {
            margin: 8px 0;
            padding: 8px;
            background: white;
            border: 1px solid #dee2e6;
            border-radius: 3px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            flex-wrap: wrap;
            font-size: 0.95em;
        }
        
        .pair-source {
            font-weight: 500;
            color: #2c3e50;
            flex: 1;
            min-width: 150px;
        }
        
        .pair-target {
            color: #7f8c8d;
            flex: 0 1 35%;
            text-align: right;
            margin-left: 10px;
            font-size: 0.9em;
        }
        
        .categorization-item {
            margin: 12px 0;
            padding: 12px;
            background: white;
            border: 1px solid #dee2e6;
            border-radius: 4px;
            break-inside: avoid;
        }
        
        .categorization-item img {
            max-width: 100%;
            max-height: 150px;
            margin: 10px 0;
            border-radius: 3px;
        }
        
        .categorization-text {
            font-size: 0.95em;
            margin-bottom: 8px;
            color: #2c3e50;
        }
        
        .categorization-input {
            border-bottom: 1.5px solid #2c3e50;
            display: inline-block;
            min-width: 100px;
            margin-top: 8px;
            height: 18px;
        }
        
        .category-list {
            margin-top: 10px;
            padding: 8px 12px;
            background: #ecf0f1;
            border-radius: 3px;
            font-size: 0.85em;
            color: #7f8c8d;
            break-inside: avoid;
        }
        
        .full-page-images {
            margin-top: 50px;
            padding-top: 20px;
            border-top: 3px dashed #95a5a6;
            break-before: page;
        }
        
        .full-page-images-header {
            background: #9b59b6;
            color: white;
            padding: 12px;
            border-radius: 4px;
            font-weight: bold;
            margin-bottom: 20px;
            font-size: 1.1em;
        }
        
        .full-page-image-item {
            margin-bottom: 30px;
            page-break-inside: avoid;
        }
        
        .full-page-image-label {
            font-weight: bold;
            color: #2c3e50;
            margin-bottom: 8px;
            font-size: 0.95em;
        }
        
        .full-page-image-item img {
            width: 100%;
            height: auto;
            border: 1px solid #dee2e6;
            border-radius: 4px;
        }
        
        .answers-section {
            margin-top: 50px;
            padding-top: 20px;
            border-top: 3px dashed #95a5a6;
            break-before: page;
        }
        
        .answers-header {
            background: #e74c3c;
            color: white;
            padding: 12px;
            border-radius: 4px;
            font-weight: bold;
            margin-bottom: 20px;
            font-size: 1.1em;
        }
        
        .answer-item {
            margin-bottom: 15px;
            padding: 12px;
            background: #f8f9fa;
            border-left: 4px solid #27ae60;
            border-radius: 3px;
            break-inside: avoid;
        }
        
        .answer-number {
            font-weight: bold;
            color: #27ae60;
            margin-bottom: 4px;
            font-size: 0.9em;
        }
        
        .answer-text {
            color: #2c3e50;
            font-size: 0.9em;
            word-break: break-word;
        }
        
        .footer {
            text-align: center;
            color: #95a5a6;
            font-size: 0.8em;
            margin-top: 40px;
            padding-top: 15px;
            border-top: 1px solid #ecf0f1;
        }
        
        @media print {
            body {
                background: white;
                padding: 0;
            }
            .container {
                max-width: 100%;
                margin: 0;
            }
            .exercise {
                break-inside: avoid;
            }
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>📚 WifeyMOOC Worksheet</h1>
            <p>Exercise Set • Generated on {date}</p>
        </div>
"""
        
        # First pass: collect full-page images
        self._collect_full_page_images()
        
        # Add exercises
        for idx, exercise in enumerate(self.exercises, 1):
            html_content += self._exercise_to_html(exercise, idx)
        
        # Add full-page images section if there are any
        if self.full_page_images:
            html_content += self._generate_full_page_images_section()
        
        # Add answers section
        html_content += self._generate_answers_section()
        
        # Close HTML
        html_content += """
        <div class="footer">
            <p>© WifeyMOOC - Keep Learning! 💫</p>
        </div>
    </div>
</body>
</html>"""
        
        html_content = html_content.replace("{date}", datetime.now().strftime("%B %d, %Y"))
        
        # Save to file
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        print(f"✓ HTML worksheet saved to: {output_file}")
        return output_file
    
    def generate_docx(self, output_file: str = None) -> str:
        """Generate DOCX document for editing"""
        if not HAS_DOCX:
            print("⚠️  Warning: python-docx not installed. Install with: pip install python-docx")
            return None
        
        if output_file is None:
            output_file = Path(self.json_file).stem + "_paper.docx"
        
        doc = Document()
        
        # Set up A4 page size
        section = doc.sections[0]
        section.page_height = Inches(11.7)
        section.page_width = Inches(8.3)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        # Add header
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("📚 WifeyMOOC Worksheet")
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(44, 62, 80)
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f"Exercise Set • Generated on {datetime.now().strftime('%B %d, %Y')}")
        subtitle_run.font.size = Pt(11)
        subtitle_run.font.color.rgb = RGBColor(127, 140, 141)
        
        doc.add_paragraph()  # Spacing
        
        # First pass: collect full-page images
        self._collect_full_page_images()
        
        # Add exercises
        for idx, exercise in enumerate(self.exercises, 1):
            self._exercise_to_docx(doc, exercise, idx)
        
        # Add page break before reference images
        if self.full_page_images:
            doc.add_page_break()
            self._generate_full_page_images_docx(doc)
        
        # Add page break before answer key
        doc.add_page_break()
        self._generate_answers_docx(doc)
        
        # Save to file
        doc.save(output_file)
        print(f"✓ DOCX worksheet saved to: {output_file}")
        return output_file
    
    def _collect_full_page_images(self):
        """Collect images from simple types (MCQ, list_pick, etc.) for full-page rendering"""
        self.full_page_images = {}  # Reset
        for idx, exercise in enumerate(self.exercises, 1):
            ex_type = exercise.get('type', '')
            # Only collect images from simple question types
            if ex_type in ['mcq_single', 'mcq_multiple', 'list_pick', 'fill_blanks_dropdown', 
                          'word_fill', 'sequence_audio', 'order_phrase']:
                if exercise.get('media'):
                    media = exercise['media']
                    if isinstance(media, dict) and 'image' in media:
                        image_path = media['image']
                        question_num = idx
                        question_text = exercise.get('question', 'N/A')
                        self.full_page_images[image_path] = {
                            'question_num': question_num,
                            'question_text': question_text,
                            'type': ex_type
                        }
    
    def _exercise_to_html(self, exercise: Dict[str, Any], number: int) -> str:
        """Convert individual exercise to HTML"""
        ex_type = exercise.get('type', 'unknown')
        html = f'''
        <div class="exercise">
            <div>
                <span class="exercise-number">Q{number}</span>
                <span class="exercise-type">{ex_type.upper()}</span>
            </div>
            <div class="question">{exercise.get('question', 'N/A')}</div>
'''
        
        # Handle media if present
        if exercise.get('media'):
            media = exercise['media']
            if isinstance(media, dict):
                if 'video' in media:
                    html += f'<div class="media-note">🎥 Video: {media["video"]}</div>\n'
                if 'audio' in media:
                    html += f'<div class="media-note">🔊 Audio: {media["audio"]}</div>\n'
                if 'image' in media:
                    # Check if this is a simple type (MCQ, etc.) that should reference full page
                    if ex_type in ['mcq_single', 'mcq_multiple', 'list_pick', 'fill_blanks_dropdown',
                                  'word_fill', 'sequence_audio', 'order_phrase']:
                        # Reference full-page image instead of embedding
                        image_filename = Path(media['image']).name
                        html += f'<div class="image-reference">📄 See image "{image_filename}" on the reference page</div>\n'
                    elif ex_type not in ['match_sentence', 'image_tagging']:
                        # For other types (not match_sentence or image_tagging), embed inline
                        html += self._render_inline_image(media['image'])
        
        # Type-specific rendering
        if ex_type == 'mcq_single':
            html += self._render_mcq_single(exercise)
        elif ex_type == 'mcq_multiple':
            html += self._render_mcq_multiple(exercise)
        elif ex_type == 'list_pick':
            html += self._render_list_pick(exercise)
        elif ex_type == 'fill_blanks_dropdown':
            html += self._render_fill_blanks(exercise)
        elif ex_type == 'match_phrases':
            html += self._render_match_phrases(exercise)
        elif ex_type == 'match_sentence':
            html += self._render_match_sentence(exercise)
        elif ex_type == 'order_phrase':
            html += self._render_order_phrase(exercise)
        elif ex_type == 'categorization_multiple':
            html += self._render_categorization(exercise)
        elif ex_type == 'word_fill':
            html += self._render_word_fill(exercise)
        elif ex_type == 'sequence_audio':
            html += self._render_sequence(exercise)
        elif ex_type == 'image_tagging':
            html += self._render_image_tagging(exercise)
        elif ex_type == 'multi_questions':
            html += self._render_multi_questions(exercise)
        else:
            html += '<div class="media-note">⚠️ Exercise type not yet formatted for printing</div>\n'
        
        html += '</div>\n'
        return html
    
    def _exercise_to_docx(self, doc: Document, exercise: Dict[str, Any], number: int):
        """Convert individual exercise to DOCX"""
        ex_type = exercise.get('type', 'unknown')
        
        # Add exercise header
        header = doc.add_paragraph()
        q_num = header.add_run(f"Q{number}")
        q_num.font.size = Pt(11)
        q_num.font.bold = True
        q_num.font.color.rgb = RGBColor(52, 152, 219)
        
        q_type = header.add_run(f" [{ex_type.upper()}]")
        q_type.font.size = Pt(10)
        q_type.font.color.rgb = RGBColor(44, 62, 80)
        
        # Add question
        q_para = doc.add_paragraph(exercise.get('question', 'N/A'))
        q_para.style = 'Normal'
        q_run = q_para.runs[0]
        q_run.font.size = Pt(11)
        q_run.font.bold = True
        q_run.font.color.rgb = RGBColor(44, 62, 80)
        
        # Handle media
        if exercise.get('media'):
            media = exercise['media']
            if isinstance(media, dict):
                if 'video' in media:
                    note = doc.add_paragraph(f"🎥 Video: {media['video']}")
                    self._style_note_paragraph(note)
                if 'audio' in media:
                    note = doc.add_paragraph(f"🔊 Audio: {media['audio']}")
                    self._style_note_paragraph(note)
                if 'image' in media:
                    # Check if should reference or embed
                    if ex_type in ['mcq_single', 'mcq_multiple', 'list_pick', 'fill_blanks_dropdown',
                                  'word_fill', 'sequence_audio', 'order_phrase']:
                        # Reference full-page image
                        image_filename = Path(media['image']).name
                        note = doc.add_paragraph(f"📄 See image \"{image_filename}\" on the reference page")
                        self._style_note_paragraph(note)
                    elif ex_type not in ['match_sentence', 'image_tagging']:
                        # Embed image
                        self._add_image_to_docx(doc, media['image'])
        
        # Type-specific rendering
        if ex_type == 'mcq_single':
            self._render_mcq_single_docx(doc, exercise)
        elif ex_type == 'mcq_multiple':
            self._render_mcq_multiple_docx(doc, exercise)
        elif ex_type == 'list_pick':
            self._render_list_pick_docx(doc, exercise)
        elif ex_type == 'fill_blanks_dropdown':
            self._render_fill_blanks_docx(doc, exercise)
        elif ex_type == 'match_phrases':
            self._render_match_phrases_docx(doc, exercise)
        elif ex_type == 'match_sentence':
            self._render_match_sentence_docx(doc, exercise)
        elif ex_type == 'order_phrase':
            self._render_order_phrase_docx(doc, exercise)
        elif ex_type == 'categorization_multiple':
            self._render_categorization_docx(doc, exercise)
        elif ex_type == 'word_fill':
            self._render_word_fill_docx(doc, exercise)
        elif ex_type == 'sequence_audio':
            self._render_sequence_docx(doc, exercise)
        elif ex_type == 'image_tagging':
            self._render_image_tagging_docx(doc, exercise)
        elif ex_type == 'multi_questions':
            self._render_multi_questions_docx(doc, exercise)
        
        doc.add_paragraph()  # Spacing
    
    def _style_note_paragraph(self, para):
        """Style a note paragraph"""
        para.style = 'Normal'
        for run in para.runs:
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(133, 100, 4)
    
    def _add_image_to_docx(self, doc: Document, image_path: str, width: float = 5.0):
        """Add an image to DOCX document"""
        try:
            full_path = self._get_image_path(image_path)
            if full_path:
                doc.add_picture(str(full_path), width=Inches(width))
            else:
                doc.add_paragraph(f"[Image not found: {image_path}]")
        except Exception as e:
            doc.add_paragraph(f"[Could not add image: {image_path}]")
    
    def _add_checkbox_to_docx(self, para: Document, text: str, checked: bool = False):
        """Add a checkbox to a paragraph in DOCX"""
        checkbox_symbol = "☑" if checked else "☐"
        p = para.add_paragraph(f"{checkbox_symbol} {text}")
        return p
    
    def _render_mcq_single_docx(self, doc: Document, exercise: Dict):
        """Render MCQ single choice in DOCX with checkboxes"""
        for option in exercise.get('options', []):
            self._add_checkbox_to_docx(doc, option)
    
    def _render_mcq_multiple_docx(self, doc: Document, exercise: Dict):
        """Render MCQ multiple choice in DOCX with checkboxes"""
        for option in exercise.get('options', []):
            self._add_checkbox_to_docx(doc, option)
    
    def _render_list_pick_docx(self, doc: Document, exercise: Dict):
        """Render list pick in DOCX with checkboxes"""
        for option in exercise.get('options', []):
            self._add_checkbox_to_docx(doc, option)
    
    def _render_fill_blanks_docx(self, doc: Document, exercise: Dict):
        """Render fill blanks in DOCX with dropdown options"""
        # Show the sentence with blanks
        parts = exercise.get('sentence_parts', [])
        p = doc.add_paragraph()
        for part in parts:
            p.add_run(part)
        
        # Show available options
        options = exercise.get('options_for_blanks', [])
        if options:
            doc.add_paragraph("Available options:", style='Heading 4')
            for option in options:
                if isinstance(option, list):
                    # If option is a list of choices for that blank
                    doc.add_paragraph("Blank options:")
                    for choice in option:
                        self._add_checkbox_to_docx(doc, choice)
                    doc.add_paragraph()  # Spacing between blanks
                else:
                    self._add_checkbox_to_docx(doc, str(option))
    
    def _render_match_phrases_docx(self, doc: Document, exercise: Dict):
        """Render match phrases in DOCX"""
        doc.add_paragraph("Match the phrases:", style='Heading 4')
        pairs = exercise.get('pairs', [])
        for pair in pairs:
            source = pair.get('source', '')
            p = doc.add_paragraph()
            p.add_run(source).bold = True
            p.add_run("  " + "_" * 20)
    
    def _render_match_sentence_docx(self, doc: Document, exercise: Dict):
        """Render match sentence in DOCX"""
        pairs = exercise.get('pairs', [])
        
        # Shuffle and display images
        images_with_indices = [(idx, pair.get('image_path', '')) for idx, pair in enumerate(pairs)]
        random_order = list(range(len(pairs)))
        random.shuffle(random_order)
        shuffled_images = [images_with_indices[i] for i in random_order]
        
        doc.add_paragraph("Images:", style='Heading 4')
        for label_idx, (orig_idx, image_path) in enumerate(shuffled_images):
            full_path = self._get_image_path(image_path)
            if full_path:
                label = chr(65 + label_idx)
                p = doc.add_paragraph(f"({label}) ", style='Normal')
                try:
                    doc.add_picture(str(full_path), width=Inches(2.0))
                except:
                    pass
        
        doc.add_paragraph("Match sentences with images:", style='Heading 4')
        for idx, pair in enumerate(pairs, 1):
            sentence = pair.get('sentence', '')
            p = doc.add_paragraph()
            p.add_run(f"{idx}. {sentence}").bold = True
            p.add_run("  " + "_" * 15)
    
    def _render_order_phrase_docx(self, doc: Document, exercise: Dict):
        """Render order phrase in DOCX"""
        doc.add_paragraph("Order the sentences:", style='Heading 4')
        for idx, phrase in enumerate(exercise.get('phrase_shuffled', []), 1):
            p = doc.add_paragraph()
            p.add_run("__ ").bold = True
            p.add_run(f"{idx}. {phrase}")
    
    def _render_categorization_docx(self, doc: Document, exercise: Dict):
        """Render categorization in DOCX"""
        categories = exercise.get('categories', [])
        cat_text = ', '.join(c for c in categories if c.strip())
        doc.add_paragraph(f"Categories: {cat_text}", style='Heading 4')
        
        stimuli = exercise.get('stimuli', [])
        for stimulus in stimuli:
            text = stimulus.get('text', '')
            image = stimulus.get('image', '')
            
            if text:
                p = doc.add_paragraph(text)
            
            if image:
                self._add_image_to_docx(doc, image, width=3.0)
            
            p = doc.add_paragraph()
            p.add_run("Category: ").bold = True
            p.add_run("_" * 15)
            doc.add_paragraph()  # Spacing
    
    def _render_word_fill_docx(self, doc: Document, exercise: Dict):
        """Render word fill in DOCX"""
        parts = exercise.get('sentence_parts', [])
        p = doc.add_paragraph()
        for part in parts:
            p.add_run(part)
        
        for idx in range(len(exercise.get('answers', []))):
            blank = doc.add_paragraph()
            blank.add_run(f"Answer {idx+1}: ").bold = True
            blank.add_run("_" * 20)
    
    def _render_sequence_docx(self, doc: Document, exercise: Dict):
        """Render sequence in DOCX"""
        doc.add_paragraph("Put items in order:", style='Heading 4')
        options = exercise.get('audio_options', [])
        for idx, option in enumerate(options, 1):
            opt_text = option.get('option', f'Item {idx}') if isinstance(option, dict) else option
            p = doc.add_paragraph()
            p.add_run("__ ").bold = True
            p.add_run(opt_text)
    
    def _render_image_tagging_docx(self, doc: Document, exercise: Dict):
        """Render image tagging in DOCX"""
        media = exercise.get('media', {})
        image_path = media.get('image', '')
        
        if image_path:
            self._add_image_to_docx(doc, image_path, width=5.0)
        
        doc.add_paragraph(f"Button: {exercise.get('button_label', 'N/A')}", style='Heading 4')
        doc.add_paragraph("Label the diagram with:")
        
        tags = exercise.get('tags', [])
        for tag in tags:
            doc.add_paragraph(tag.get('label', 'N/A'), style='List Bullet')
    
    def _render_multi_questions_docx(self, doc: Document, exercise: Dict):
        """Render multi questions in DOCX"""
        doc.add_paragraph("Sub-questions:", style='Heading 4')
        for q_idx, question in enumerate(exercise.get('questions', []), 1):
            p = doc.add_paragraph()
            p.add_run(f"Q{q_idx}: ").bold = True
            p.add_run(question.get('question', 'N/A'))
    
    def _generate_full_page_images_section(self) -> str:
        """Generate full-page images section for HTML"""
        html = '<div class="full-page-images">\n'
        html += '<div class="full-page-images-header">🖼️ REFERENCE IMAGES</div>\n'
        
        for image_path, metadata in self.full_page_images.items():
            data_uri = self._load_image_as_base64(image_path)
            if data_uri:
                question_num = metadata['question_num']
                question_text = metadata['question_text']
                image_filename = Path(image_path).name
                
                html += f'<div class="full-page-image-item">\n'
                html += f'<div class="full-page-image-label">Question {question_num}: {image_filename}</div>\n'
                html += f'<p style="margin-bottom: 10px; color: #7f8c8d; font-size: 0.9em; font-style: italic;">{question_text}</p>\n'
                html += f'<img src="{data_uri}" alt="Question {question_num} image">\n'
                html += '</div>\n'
        
        html += '</div>\n'
        return html
    
    def _generate_full_page_images_docx(self, doc: Document):
        """Generate full-page images section for DOCX"""
        title = doc.add_paragraph()
        title_run = title.add_run("🖼️ REFERENCE IMAGES")
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(155, 89, 182)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        for image_path, metadata in self.full_page_images.items():
            full_path = self._get_image_path(image_path)
            if full_path:
                question_num = metadata['question_num']
                question_text = metadata['question_text']
                image_filename = Path(image_path).name
                
                # Title
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(f"Question {question_num}: {image_filename}")
                title_run.font.bold = True
                title_run.font.size = Pt(11)
                
                # Question text
                q_para = doc.add_paragraph(question_text)
                for run in q_para.runs:
                    run.font.italic = True
                    run.font.size = Pt(10)
                
                # Image
                try:
                    doc.add_picture(str(full_path), width=Inches(6.0))
                except:
                    doc.add_paragraph(f"[Could not load image: {image_filename}]")
                
                doc.add_paragraph()  # Spacing
    
    def _generate_answers_section(self) -> str:
        """Generate answers section for HTML"""
        html = '<div class="answers-section">\n'
        html += '<div class="answers-header">📋 ANSWER KEY</div>\n'
        
        for idx, exercise in enumerate(self.exercises, 1):
            answer = exercise.get('answer')
            if answer is None:
                continue
            
            html += f'<div class="answer-item">\n'
            html += f'<div class="answer-number">Question {idx}</div>\n'
            
            if isinstance(answer, list):
                answer_text = ', '.join(str(a) for a in answer) if answer else 'N/A'
            elif isinstance(answer, dict):
                answer_text = ' | '.join([f'{k}: {v}' for k, v in answer.items()])
            else:
                answer_text = str(answer)
            
            html += f'<div class="answer-text">{answer_text}</div>\n'
            html += '</div>\n'
        
        html += '</div>\n'
        return html
    
    def _generate_answers_docx(self, doc: Document):
        """Generate answers section for DOCX"""
        title = doc.add_paragraph()
        title_run = title.add_run("📋 ANSWER KEY")
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(231, 76, 60)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        for idx, exercise in enumerate(self.exercises, 1):
            answer = exercise.get('answer')
            if answer is None:
                continue
            
            # Question number
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"Q{idx}")
            q_run.font.bold = True
            q_run.font.color.rgb = RGBColor(39, 174, 96)
            
            # Answer
            if isinstance(answer, list):
                answer_text = ', '.join(str(a) for a in answer) if answer else 'N/A'
            elif isinstance(answer, dict):
                answer_text = ' | '.join([f'{k}: {v}' for k, v in answer.items()])
            else:
                answer_text = str(answer)
            
            a_run = q_para.add_run(f": {answer_text}")
            a_run.font.size = Pt(11)
    
    def _render_inline_image(self, image_path: str) -> str:
        """Render an image inline with base64 encoding for HTML"""
        data_uri = self._load_image_as_base64(image_path)
        if data_uri:
            filename = Path(image_path).name
            return f'''<div class="media-image">
    <img src="{data_uri}" alt="{filename}">
    <div class="media-image-caption">{filename}</div>
</div>\n'''
        else:
            return f'<div class="media-note">🖼️ Image: {image_path}</div>\n'
    
    def _render_mcq_single(self, exercise: Dict) -> str:
        """Render MCQ single choice for HTML"""
        html = '<div class="option-list">\n'
        for idx, option in enumerate(exercise.get('options', [])):
            html += f'<div class="option"><input type="radio" id="q_opt_{idx}" name="question"> <label for="q_opt_{idx}">{option}</label></div>\n'
        html += '</div>\n'
        return html
    
    def _render_mcq_multiple(self, exercise: Dict) -> str:
        """Render MCQ multiple choice for HTML"""
        html = '<div class="option-list">\n'
        for idx, option in enumerate(exercise.get('options', [])):
            html += f'<div class="option"><input type="checkbox" id="q_opt_{idx}" name="question"> <label for="q_opt_{idx}">{option}</label></div>\n'
        html += '</div>\n'
        return html
    
    def _render_list_pick(self, exercise: Dict) -> str:
        """Render list pick for HTML"""
        html = '<div class="option-list">\n'
        for idx, option in enumerate(exercise.get('options', [])):
            html += f'<div class="option"><input type="checkbox" id="q_opt_{idx}" name="question"> <label for="q_opt_{idx}">{option}</label></div>\n'
        html += '</div>\n'
        return html
    
    def _render_fill_blanks(self, exercise: Dict) -> str:
        """Render fill blanks for HTML with dropdown options displayed"""
        html = '<div class="sentence-parts">\n'
        parts = exercise.get('sentence_parts', [])
        
        # Show sentence with blanks
        for part in parts:
            html += part
        html += '\n'
        
        # Show available options
        options = exercise.get('options_for_blanks', [])
        if options:
            html += '<div class="category-list"><strong>Available options:</strong> '
            option_list = []
            for option in options:
                if isinstance(option, list):
                    option_list.extend(option)
                else:
                    option_list.append(str(option))
            html += ', '.join(option_list)
            html += '</div>\n'
        
        html += '</div>\n'
        return html
    
    def _render_match_phrases(self, exercise: Dict) -> str:
        """Render match phrases for HTML"""
        html = '<div class="pairs-list">\n'
        for pair in exercise.get('pairs', []):
            source = pair.get('source', '')
            target = pair.get('target', '')
            html += f'<div class="pair"><div class="pair-source">{source}</div><div class="pair-target">{target}</div></div>\n'
        html += '</div>\n'
        return html
    
    def _render_match_sentence(self, exercise: Dict) -> str:
        """Render match sentence for HTML"""
        html = '<div class="image-grid">\n'
        pairs = exercise.get('pairs', [])
        
        # Shuffle images
        indices = list(range(len(pairs)))
        random.shuffle(indices)
        
        for label_idx, orig_idx in enumerate(indices):
            pair = pairs[orig_idx]
            image_path = pair.get('image_path', '')
            data_uri = self._load_image_as_base64(image_path)
            
            if data_uri:
                label = chr(65 + label_idx)
                html += f'<div class="image-item"><img src="{data_uri}" alt="Option {label}"><div class="image-label">({label})</div></div>\n'
        
        html += '</div>\n'
        
        # Add sentences to match
        html += '<div class="pairs-list"><strong>Match sentences with images:</strong>\n'
        for idx, pair in enumerate(pairs, 1):
            sentence = pair.get('sentence', '')
            html += f'<div class="pair"><div class="pair-source">{idx}. {sentence}</div><div class="pair-target">___</div></div>\n'
        html += '</div>\n'
        
        return html
    
    def _render_order_phrase(self, exercise: Dict) -> str:
        """Render order phrase for HTML"""
        html = '<div class="sentence-parts">\n'
        phrases = exercise.get('phrase_shuffled', [])
        
        for idx, phrase in enumerate(phrases, 1):
            html += f'<div class="option">___ {idx}. {phrase}</div>\n'
        
        html += '</div>\n'
        return html
    
    def _render_categorization(self, exercise: Dict) -> str:
        """Render categorization for HTML"""
        html = '<div>\n'
        categories = exercise.get('categories', [])
        cat_text = ', '.join(c for c in categories if c.strip())
        html += f'<div class="category-list"><strong>Categories:</strong> {cat_text}</div>\n'
        
        for stimulus in exercise.get('stimuli', []):
            html += '<div class="categorization-item">\n'
            
            if stimulus.get('text'):
                html += f'<div class="categorization-text">{stimulus["text"]}</div>\n'
            
            if stimulus.get('image'):
                data_uri = self._load_image_as_base64(stimulus['image'])
                if data_uri:
                    html += f'<img src="{data_uri}" alt="stimulus">\n'
            
            html += '<div class="categorization-input">Category: ___________</div>\n'
            html += '</div>\n'
        
        html += '</div>\n'
        return html
    
    def _render_word_fill(self, exercise: Dict) -> str:
        """Render word fill for HTML"""
        html = '<div class="sentence-parts">\n'
        parts = exercise.get('sentence_parts', [])
        
        for part in parts:
            html += part
        
        html += '\n'
        for idx in range(len(exercise.get('answers', []))):
            html += f'<div class="category-list">Blank {idx+1}: ___________</div>\n'
        
        html += '</div>\n'
        return html
    
    def _render_sequence(self, exercise: Dict) -> str:
        """Render sequence for HTML"""
        html = '<div class="sentence-parts">\n'
        
        for idx, option in enumerate(exercise.get('audio_options', []), 1):
            opt_text = option.get('option', f'Item {idx}') if isinstance(option, dict) else option
            html += f'<div class="option">___ {idx}. {opt_text}</div>\n'
        
        html += '</div>\n'
        return html
    
    def _render_image_tagging(self, exercise: Dict) -> str:
        """Render image tagging for HTML"""
        html = '<div>\n'
        
        media = exercise.get('media', {})
        if media.get('image'):
            data_uri = self._load_image_as_base64(media['image'])
            if data_uri:
                html += f'<div class="media-image"><img src="{data_uri}" alt="tagging"></div>\n'
        
        button_label = exercise.get('button_label', 'N/A')
        html += f'<div class="category-list"><strong>Button:</strong> {button_label}</div>\n'
        html += '<div class="category-list"><strong>Label with:</strong>\n'
        
        for tag in exercise.get('tags', []):
            label = tag.get('label', 'N/A')
            html += f'• {label}<br>\n'
        
        html += '</div>\n</div>\n'
        return html
    
    def _render_multi_questions(self, exercise: Dict) -> str:
        """Render multi questions for HTML"""
        html = '<div class="pairs-list">\n'
        
        for q_idx, question in enumerate(exercise.get('questions', []), 1):
            q_text = question.get('question', 'N/A')
            html += f'<div class="pair"><div class="pair-source">Q{q_idx}: {q_text}</div><div class="pair-target">_________</div></div>\n'
        
        html += '</div>\n'
        return html


def main():
    """Main function"""
    if len(sys.argv) < 2:
        print("Usage: python json_to_paper.py <json_file> [output_base_name]")
        print("\nExample:")
        print("  python json_to_paper.py testfile-complete.json")
        print("  python json_to_paper.py testfile-complete.json my_worksheet")
        print("\nGenerates: my_worksheet_paper.html and my_worksheet_paper.docx")
        sys.exit(1)
    
    json_file = sys.argv[1]
    output_base = sys.argv[2] if len(sys.argv) > 2 else None
    
    converter = ExerciseToPaper(json_file)
    
    # Generate HTML
    if output_base:
        html_output = f"{output_base}_paper.html"
    else:
        html_output = None
    html_file = converter.generate_html(html_output)
    
    # Generate DOCX
    if output_base:
        docx_output = f"{output_base}_paper.docx"
    else:
        docx_output = None
    docx_file = converter.generate_docx(docx_output)
    
    print(f"\n✨ Worksheets ready!")
    print(f"📄 HTML: {html_file}")
    if docx_file:
        print(f"📝 DOCX: {docx_file}")
        print(f"\n💡 Use DOCX for editing, HTML for printing to PDF")
    else:
        print(f"⚠️  Install python-docx for DOCX support: pip install python-docx")


if __name__ == "__main__":
    main()
